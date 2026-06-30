# utils/data_loader.py
import os
import requests
import zipfile
import io
from pathlib import Path
from typing import List, Dict, Optional, Union
import logging
import numpy as np
from tqdm import tqdm # Ajout de tqdm

import re
import unicodedata

import logfire
from pydantic import ValidationError

from .schemas import CleanedDocument, DocumentMetadata, RawDocument

_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_MULTI_SPACES_RE = re.compile(r"[ \t]{2,}")
_MULTI_BLANK_LINES_RE = re.compile(r"\n{3,}")


def clean_text(text: str) -> str:
    """Normalise le texte extrait avant chunking (espaces, BOM, caractères de contrôle)."""
    text = text.replace("﻿", "")
    text = unicodedata.normalize("NFKC", text)
    text = _CONTROL_CHARS_RE.sub("", text)
    text = _MULTI_SPACES_RE.sub(" ", text)
    text = _MULTI_BLANK_LINES_RE.sub("\n\n", text)
    return text.strip()


def build_cleaned_document(page_content: str, metadata: dict) -> CleanedDocument | None:
    """Valide le contenu brut, le nettoie, puis valide le résultat final.

    Retourne None (et logue une erreur) si le document est vide ou invalide
    à n'importe quelle étape, plutôt que de faire échouer tout le pipeline.
    """
    try:
        raw = RawDocument(page_content=page_content, metadata=DocumentMetadata(**metadata))
    except ValidationError as e:
        logging.error(f"Document brut invalide ignoré ({metadata.get('source', '?')}): {e}")
        return None

    cleaned_text = clean_text(raw.page_content)
    try:
        return CleanedDocument(page_content=cleaned_text, metadata=raw.metadata)
    except ValidationError as e:
        logging.error(f"Document nettoyé invalide ignoré ({metadata.get('source', '?')}): {e}")
        return None

# --- Importations pour OCR ---
try:
    import fitz  # PyMuPDF
    from PIL import Image
    import easyocr

    # Initialiser le lecteur EasyOCR une seule fois
    logging.info("Initialisation du lecteur EasyOCR...")
    reader = easyocr.Reader(['en', 'fr']) 
    logging.info("Lecteur EasyOCR initialisé.")

except ImportError as e:
    logging.warning(f"Modules OCR (PyMuPDF, Pillow, easyocr) non installés ou erreur: {e}. L'OCR pour PDF ne sera pas disponible.")
    fitz = None
    Image = None
    easyocr = None
    reader = None
except Exception as e:
    logging.error(f"Erreur inattendue lors du chargement des modules/modèle OCR: {e}")
    fitz = None
    Image = None
    easyocr = None
    reader = None

# Configuration du logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Fonctions d'extraction de texte ---

def extract_text_from_pdf_with_ocr(file_path: str) -> Optional[str]:
    """Extrait le texte d'un fichier PDF en utilisant l'OCR (EasyOCR)."""
    if not fitz or not reader:
        logging.warning("Modules/Modèle OCR non disponibles. Impossible d'effectuer l'OCR.")
        return None

    text_content = []
    try:
        doc = fitz.open(file_path)
        # Utiliser tqdm pour la barre de progression
        for page_num in tqdm(range(len(doc)), desc=f"OCR de {os.path.basename(file_path)}"):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # Augmenter la résolution pour l'OCR
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            try:
                img_np = np.array(img)
                results = reader.readtext(img_np)
                page_text = "\n".join([res[1] for res in results])
                text_content.append(page_text)
                # logging.info(f"OCR effectuée sur la page {page_num + 1} de {file_path} avec EasyOCR") # Commenté pour éviter le spam de logs avec tqdm
            except Exception as ocr_e:
                logging.error(f"Erreur lors de l'OCR de la page {page_num + 1} de {file_path} avec EasyOCR: {ocr_e}")
                continue

        doc.close()
        full_text = "\n".join(text_content).strip()
        if full_text:
            logging.info(f"Texte extrait via OCR de PDF: {file_path} ({len(full_text)} caractères)")
            return full_text
        else:
            logging.warning(f"Aucun texte significatif extrait via OCR de {file_path}.")
            return None
    except Exception as e:
        logging.error(f"Erreur lors de l'ouverture ou du traitement OCR du PDF {file_path}: {e}")
        return None

def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """Extrait le texte d'un fichier PDF, avec fallback OCR si peu de texte est trouvé."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = "".join(page.extract_text() + "\n" for page in reader.pages if page.extract_text())
        
        if len(text.strip()) < 100: # Si très peu de texte est extrait, tenter l'OCR
            logging.info(f"Peu de texte trouvé dans {file_path} via extraction standard ({len(text.strip())} caractères). Tentative d'OCR...")
            ocr_text = extract_text_from_pdf_with_ocr(file_path)
            if ocr_text:
                return ocr_text
            else:
                logging.warning(f"L'OCR n'a pas non plus produit de texte significatif pour {file_path}.")
                return text # Retourne le peu de texte trouvé ou vide
        
        logging.info(f"Texte extrait de PDF: {file_path} ({len(text)} caractères)")
        return text
    except Exception as e:
        logging.error(f"Erreur extraction PDF {file_path}: {e}. Tentative d'OCR en dernier recours...")
        # Si l'extraction standard échoue complètement, tenter l'OCR
        ocr_text = extract_text_from_pdf_with_ocr(file_path)
        if ocr_text:
            return ocr_text
        else:
            logging.warning(f"L'OCR n'a pas non plus produit de texte significatif après échec de l'extraction standard pour {file_path}.")
            return None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """Extrait le texte d'un fichier Word DOCX."""
    try:
        import docx
        doc = docx.Document(file_path)
        text = "\n".join(para.text for para in doc.paragraphs if para.text)
        logging.info(f"Texte extrait de DOCX: {file_path} ({len(text)} caractères)")
        return text
    except Exception as e:
        logging.error(f"Erreur extraction DOCX {file_path}: {e}")
        return None

def extract_text_from_txt(file_path: str) -> Optional[str]:
    """Extrait le texte d'un fichier texte brut."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        logging.info(f"Texte extrait de TXT: {file_path} ({len(text)} caractères)")
        return text
    except Exception as e:
        logging.error(f"Erreur extraction TXT {file_path}: {e}")
        return None

def extract_text_from_csv(file_path: str) -> Optional[str]:
    """Extrait le texte d'un fichier CSV (convertit en string)."""
    try:
        import pandas as pd
        try:
            df = pd.read_csv(file_path)
        except UnicodeDecodeError:
            df = pd.read_csv(file_path, encoding='latin1') # Essayer un autre encodage courant
        except Exception as read_e:
             logging.warning(f"Erreur lecture CSV {file_path}: {read_e}. Tentative avec séparateur ';'")
             try:
                 df = pd.read_csv(file_path, sep=';')
             except UnicodeDecodeError:
                  df = pd.read_csv(file_path, sep=';', encoding='latin1')
             except Exception as read_e2:
                   logging.error(f"Impossible de lire le CSV {file_path}: {read_e2}")
                   return None

        text = df.to_string()
        logging.info(f"Texte extrait de CSV: {file_path} ({len(text)} caractères)")
        return text
    except ImportError:
        logging.warning("Pandas non installé. Impossible de lire les fichiers CSV.")
        return None
    except Exception as e:
        logging.error(f"Erreur extraction CSV {file_path}: {e}")
        return None

def extract_text_from_excel(file_path: str) -> Optional[Union[str, Dict[str, str]]]:
    """Extrait le texte de chaque feuille d'un fichier Excel."""
    try:
        import pandas as pd
        # Lire toutes les feuilles dans un dictionnaire de DataFrames
        excel_file = pd.ExcelFile(file_path)
        sheets_data = {}
        for sheet_name in excel_file.sheet_names:
            df = excel_file.parse(sheet_name)
            sheets_data[sheet_name] = df.to_string()
        
        logging.info(f"Texte extrait de {len(sheets_data)} feuille(s) dans Excel: {file_path}")
        # Si une seule feuille, retourne directement le texte pour la compatibilité
        if len(sheets_data) == 1:
            return list(sheets_data.values())[0]
        return sheets_data
    except ImportError:
        logging.warning("Pandas ou openpyxl non installé. Impossible de lire les fichiers Excel.")
        return None
    except Exception as e:
        logging.error(f"Erreur extraction Excel {file_path}: {e}")
        return None

# --- Fonctions de chargement ---

def download_and_extract_zip(url: str, output_dir: str) -> bool:
    """Télécharge un fichier ZIP depuis une URL et l'extrait."""
    if not url:
        logging.warning("Aucune URL fournie pour le téléchargement.")
        return False
    try:
        logging.info(f"Téléchargement des données depuis {url}...")
        response = requests.get(url, stream=True)
        response.raise_for_status()

        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        with zipfile.ZipFile(io.BytesIO(response.content)) as z:
            logging.info(f"Extraction du contenu dans {output_dir}...")
            z.extractall(output_dir)
        logging.info("Téléchargement et extraction terminés.")
        return True
    except requests.exceptions.RequestException as e:
        logging.error(f"Erreur de téléchargement: {e}")
        return False
    except zipfile.BadZipFile:
        logging.error("Le fichier téléchargé n'est pas un ZIP valide.")
        return False
    except Exception as e:
        logging.error(f"Erreur inattendue lors du téléchargement/extraction: {e}")
        return False

def load_and_parse_files(input_dir: str) -> List[CleanedDocument]:
    """
    Charge et parse récursivement les fichiers d'un répertoire.
    Retourne une liste de dictionnaires, chacun représentant un document.
    """
    documents: List[CleanedDocument] = []
    input_path = Path(input_dir)
    if not input_path.is_dir():
        logging.error(f"Le répertoire d'entrée '{input_dir}' n'existe pas.")
        return []

    logging.info(f"Parcours du répertoire source: {input_dir}")
    for file_path in input_path.rglob("*.*"):
        if file_path.is_file():
            relative_path = file_path.relative_to(input_path)
            source_folder = relative_path.parts[0] if len(relative_path.parts) > 1 else "root"
            ext = file_path.suffix.lower()
            
            logging.debug(f"Traitement du fichier: {relative_path} (Dossier source: {source_folder})")

            with logfire.span("extract_and_clean_document", file=str(relative_path)):
                extracted_content = None
                if ext == ".pdf":
                    extracted_content = extract_text_from_pdf(str(file_path))
                elif ext == ".docx":
                    extracted_content = extract_text_from_docx(str(file_path))
                elif ext == ".txt":
                    extracted_content = extract_text_from_txt(str(file_path))
                elif ext == ".csv":
                    extracted_content = extract_text_from_csv(str(file_path))
                elif ext in [".xlsx", ".xls"]:
                    extracted_content = extract_text_from_excel(str(file_path))
                # Suppression de la gestion des fichiers HTML
                else:
                    logging.warning(f"Type de fichier non supporté ignoré: {relative_path}")
                    continue

                if not extracted_content:
                    logging.warning(f"Aucun contenu n'a pu être extrait de {relative_path}")
                    continue

                # Si c'est un dictionnaire (plusieurs feuilles Excel), créer un doc par feuille
                if isinstance(extracted_content, dict):
                    for sheet_name, text in extracted_content.items():
                        cleaned = build_cleaned_document(text, {
                            "source": f"{str(relative_path)} (Feuille: {sheet_name})",
                            "filename": file_path.name,
                            "sheet": sheet_name,
                            "category": source_folder,
                            "full_path": str(file_path.resolve())
                        })
                        if cleaned is not None:
                            documents.append(cleaned)
                else: # Pour tous les autres types de fichiers
                    cleaned = build_cleaned_document(extracted_content, {
                        "source": str(relative_path),
                        "filename": file_path.name,
                        "category": source_folder,
                        "full_path": str(file_path.resolve())
                    })
                    if cleaned is not None:
                        documents.append(cleaned)

    logging.info(f"{len(documents)} documents chargés et parsés.")
    return documents